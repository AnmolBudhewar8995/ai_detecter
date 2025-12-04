import argparse
import re
import sys
from functools import lru_cache
from io import BytesIO
from pathlib import Path

from docx import Document
from transformers import pipeline

@lru_cache(maxsize=1)
def _load_detector_pipe():
    print("Loading Model... (First time it will take time)")
    return pipeline("text-classification", model="roberta-base-openai-detector")


def analyze_text(text: str) -> dict:
    pipe = _load_detector_pipe()
    result = pipe(text, truncation=True, max_length=512)
    label = result[0]['label']  # 'Real' (Human) किंवा 'Fake' (AI)
    score = result[0]['score'] * 100
    return {
        'label': label,
        'score': score,
        'originality': 100.0 - score,
        'text': text,
        'sentences': analyze_sentences(text),
        'paragraphs': _paragraphs_from_text(text),
    }


def _split_sentences(text: str) -> list[str]:
    sentences = [s.strip() for s in re.split(r'(?<=[.!?])\s+', text) if s.strip()]
    return sentences


def analyze_sentences(text: str) -> list[dict]:
    sentences = _split_sentences(text)
    if not sentences:
        return []

    pipe = _load_detector_pipe()
    results = pipe(sentences, truncation=True, max_length=512)
    return [
        {
            'sentence': sentences[idx],
            'label': result['label'],
            'score': result['score'] * 100,
            'originality': 100.0 - (result['score'] * 100),
        }
        for idx, result in enumerate(results)
    ]


def _paragraphs_from_text(text: str) -> list[str]:
    return [p.strip() for p in text.split('\n\n') if p.strip()]


def detect_ai_content(text: str) -> dict:
    analysis = analyze_text(text)
    score = analysis['score']
    label = analysis['label']

    print("\n--- Analysis Result ---")
    if label == 'Fake':
        print("🔴 ALERT: This text is likely AI Generated!")
        print(f"Confidence Score: {score:.2f}%")

        if score > 98:
            print("Possible Source: High probability of GPT-4 or Claude (Very structured)")
        elif score > 90:
            print("Possible Source: ChatGPT (GPT-3.5) or Gemini")
        else:
            print("Possible Source: Basic AI tool or Paraphrasing tool")
    else:
        print("🟢 This text looks Human Written.")
        print(f"Confidence Score: {score:.2f}%")

    return analysis


def _text_from_document(document: Document) -> str:
    paragraphs = [p.text for p in document.paragraphs if p.text.strip()]
    return '\n'.join(paragraphs)


def read_docx_file(path: Path) -> str:
    document = Document(path)
    return _text_from_document(document)


def read_docx_bytes(contents: bytes) -> str:
    document = Document(BytesIO(contents))
    return _text_from_document(document)


def main():
    parser = argparse.ArgumentParser(
        description="Check whether a piece of text is likely AI-generated.")
    group = parser.add_mutually_exclusive_group()
    group.add_argument(
        '-t', '--text',
        help='Text to analyze directly. Cannot be used together with --doc.',
    )
    group.add_argument(
        '-d', '--doc',
        help='Path to a Word (.docx) document. The entire text of the file will be classified.',
        type=Path,
    )
    args = parser.parse_args()

    if args.doc:
        if not args.doc.exists():
            print(f"Document not found: {args.doc}")
            return
        if args.doc.suffix.lower() != '.docx':
            print("Only .docx files are supported.")
            return
        text = read_docx_file(args.doc)
    elif args.text:
        text = args.text.strip()
    else:
        if sys.stdin.isatty():
            text = input("Enter text to analyze: ").strip()
            while not text:
                text = input("Text cannot be empty. Try again: ").strip()
        else:
            text = sys.stdin.read().strip()

    if not text:
        print("No text provided. Exiting.")
        return

    detect_ai_content(text)


if __name__ == "__main__":
    main()